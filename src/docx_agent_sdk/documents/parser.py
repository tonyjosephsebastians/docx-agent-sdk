from __future__ import annotations

from pathlib import Path

from docx import Document

from docx_agent_sdk.schemas import ExtractOptions, ExtractedDocument, ExtractedParagraph, ExtractedTable


class DocxParser:
    """Repository-style DOCX reader that normalizes document content."""

    def parse(self, path: str | Path, options: ExtractOptions | None = None) -> ExtractedDocument:
        options = options or ExtractOptions()
        document_path = Path(path)
        if not document_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {document_path}")
        if document_path.suffix.lower() != ".docx":
            raise ValueError(f"Expected a .docx file: {document_path}")

        document = Document(document_path)
        paragraphs = [
            ExtractedParagraph(index=index, text=text)
            for index, paragraph in enumerate(document.paragraphs)
            if (text := paragraph.text.strip())
        ]

        tables: list[ExtractedTable] = []
        if options.include_tables:
            for table_index, table in enumerate(document.tables):
                rows: list[list[str]] = []
                for row in table.rows:
                    rows.append([cell.text.strip() for cell in row.cells])
                tables.append(ExtractedTable(index=table_index, rows=rows))

        return ExtractedDocument(
            path=str(document_path),
            paragraphs=paragraphs,
            tables=tables,
        )
