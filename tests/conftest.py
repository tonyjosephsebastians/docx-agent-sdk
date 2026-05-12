from __future__ import annotations

from collections.abc import Sequence
from pathlib import Path

import pytest
from docx import Document

from docx_agent_sdk.providers import ModelProvider
from docx_agent_sdk.schemas import ChatMessage, GenerationOptions, GenerationResult


class FakeProvider(ModelProvider):
    def __init__(self, content: str = "fake response") -> None:
        self.content = content
        self.calls: list[list[ChatMessage]] = []

    def generate(
        self,
        messages: Sequence[ChatMessage],
        options: GenerationOptions | None = None,
    ) -> GenerationResult:
        self.calls.append(list(messages))
        return GenerationResult(content=self.content, model="fake-model")


@pytest.fixture()
def fake_provider() -> FakeProvider:
    return FakeProvider()


@pytest.fixture()
def make_docx(tmp_path: Path):
    def _make_docx(name: str, paragraphs: list[str], table: list[list[str]] | None = None) -> Path:
        path = tmp_path / name
        document = Document()
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)
        if table:
            doc_table = document.add_table(rows=len(table), cols=len(table[0]))
            for row_index, row in enumerate(table):
                for col_index, value in enumerate(row):
                    doc_table.cell(row_index, col_index).text = value
        document.save(path)
        return path

    return _make_docx
